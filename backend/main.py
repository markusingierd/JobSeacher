#!/usr/bin/env python3
"""
FinnJobScout FastAPI Backend
----------------------------
Et REST API bygget for å koble FINN-skraperen, match-analysen og AI-søknadsmotoren
sammen med TypeScript React frontend.
"""

import sys
import os
import json
import subprocess
from pathlib import Path
from typing import Optional, List, Dict, Any

from fastapi import FastAPI, HTTPException, BackgroundTasks, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response
from pydantic import BaseModel

# Legg til prosjektets rotmappe og backend-mappe i Python-stien
BASE_DIR = Path(__file__).parent.parent.resolve()
BACKEND_DIR = Path(__file__).parent.resolve()
sys.path.append(str(BASE_DIR))
sys.path.append(str(BACKEND_DIR))

import finn_scout
import job_analyst
import profile_manager
import ai_generator
import sync_upstream

app = FastAPI(
    title="FinnJobScout API",
    description="REST API for Universal FINN Job Scout & AI Cover Letter Generator",
    version="2.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

DB_FILE = BASE_DIR / "relevante_stillinger_database.json"
SOKNADSBREV_DIR = BASE_DIR / "soknadsbrev"

SOKNADSBREV_DIR.mkdir(exist_ok=True)

# -----------------------------------------------------------------------------
# PYDANTIC MODELLER
# -----------------------------------------------------------------------------

class JobModel(BaseModel):
    id: str
    title: str
    company: str
    location: str
    url: str
    status: str
    reason: Optional[str] = ""
    date_found: Optional[str] = ""
    date_published: Optional[str] = ""
    application_deadline: Optional[str] = ""
    match_percentage: int = 0
    match_analysis: Optional[str] = ""
    company_hook_insight: Optional[str] = ""
    application_status: Optional[str] = "not_applied"
    applied_file: Optional[str] = ""
    description_text: Optional[str] = ""

class GenerateApplicationRequest(BaseModel):
    job_id: str
    custom_notes: Optional[str] = ""

class ApplicationResponse(BaseModel):
    job_id: str
    job_title: str
    company: str
    cover_letter_markdown: str

class ExportDocxRequest(BaseModel):
    job_id: str
    company: str
    job_title: str
    content_markdown: str

class StatusUpdateRequest(BaseModel):
    status: str

class ProfileModel(BaseModel):
    full_name: str
    age: int
    current_title: str
    target_category: str
    tone_of_voice: str
    personality_traits: List[str]
    cv_experiences: str
    custom_search_keywords: List[str]
    location_filter: str

# -----------------------------------------------------------------------------
# HJELPEFUNKSJONER
# -----------------------------------------------------------------------------

def read_db() -> Dict[str, Any]:
    if DB_FILE.exists():
        try:
            with open(DB_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            print(f"[!] Feil ved lesing av DB: {e}")
    return {}

# -----------------------------------------------------------------------------
# API ROUTER
# -----------------------------------------------------------------------------

@app.get("/")
def read_root():
    return {
        "status": "online",
        "message": "FinnJobScout API v2.0 kjører!",
        "endpoints": [
            "/api/jobs",
            "/api/jobs/scan",
            "/api/jobs/generate-application",
            "/api/jobs/export-docx",
            "/api/profile",
            "/api/sync/check",
            "/api/sync/pull"
        ]
    }

@app.get("/api/jobs", response_model=List[JobModel])
def get_jobs(
    q: Optional[str] = Query(None, description="Søkeord (tittel, bedrift eller nøkkelord)"),
    min_match: Optional[int] = Query(0, description="Minimum match-prosent"),
    status: Optional[str] = Query(None, description="Filtrer på status (new, analyzed, excluded)"),
    app_status: Optional[str] = Query(None, description="Filtrer på søknadsstatus (not_applied, applied)")
):
    db = read_db()
    jobs = list(db.values())
    job_analyst.check_applied_status(db)
    
    filtered_jobs = []
    for job in jobs:
        if job.get("match_percentage", 0) < min_match:
            continue
        if status and job.get("status") != status:
            continue
        if app_status and job.get("application_status") != app_status:
            continue
        if q:
            q_low = q.lower()
            text_to_check = f"{job.get('title', '')} {job.get('company', '')} {job.get('match_analysis', '')} {job.get('description_text', '')}".lower()
            if q_low not in text_to_check:
                continue

        filtered_jobs.append(job)

    filtered_jobs.sort(key=lambda x: x.get("match_percentage", 0), reverse=True)
    return filtered_jobs

@app.get("/api/jobs/{job_id}", response_model=JobModel)
def get_job_by_id(job_id: str):
    db = read_db()
    if job_id not in db:
        raise HTTPException(status_code=404, detail="Stilling ikke funnet")
    return db[job_id]

@app.post("/api/jobs/scan")
def trigger_scan(background_tasks: BackgroundTasks):
    def run_pipeline():
        print("[+] Starter FINN-skanning via API...")
        finn_scout.run_scout()
        job_analyst.run_analyst()
        print("[+] FINN-skanning ferdig!")

    background_tasks.add_task(run_pipeline)
    return {"message": "FINN-skanning startet i bakgrunnen."}

@app.patch("/api/jobs/{job_id}/status")
def update_job_status(job_id: str, payload: StatusUpdateRequest):
    db = read_db()
    if job_id not in db:
        raise HTTPException(status_code=404, detail="Stilling ikke funnet")
    
    db[job_id]["application_status"] = payload.status
    job_analyst.save_database(db)
    return {"message": f"Status oppdatert til '{payload.status}'", "job": db[job_id]}

@app.get("/api/profile", response_model=ProfileModel)
def get_profile():
    return profile_manager.load_settings()

@app.post("/api/profile", response_model=ProfileModel)
def update_profile(payload: ProfileModel):
    settings = payload.model_dump()
    profile_manager.save_settings(settings)
    return settings

@app.get("/api/sync/check")
def check_sync():
    """Sjekker om kilderepositoriet har nye 1-veis oppdateringer."""
    return sync_upstream.check_upstream_updates()

@app.post("/api/sync/pull")
def pull_sync():
    """Henter nye 1-veis oppdateringer trygt fra kilderepositoriet uten å røre kilden."""
    return sync_upstream.pull_upstream_updates()

@app.post("/api/jobs/generate-application", response_model=ApplicationResponse)
def generate_application(payload: GenerateApplicationRequest):
    db = read_db()
    if payload.job_id not in db:
        raise HTTPException(status_code=404, detail="Stilling ikke funnet")

    job = db[payload.job_id]
    job_title = job.get("title", "")
    company = job.get("company", "")

    cover_letter_markdown = ai_generator.generate_cover_letter(job, payload.custom_notes or "")

    clean_company = "".join(c for c in company if c.isalnum() or c in (" ", "_")).strip().replace(" ", "_")
    draft_file = SOKNADSBREV_DIR / f"Soknad_{clean_company}_{payload.job_id}.md"
    draft_file.write_text(cover_letter_markdown, encoding="utf-8")

    db[payload.job_id]["application_status"] = "draft"
    db[payload.job_id]["applied_file"] = f"soknadsbrev/{draft_file.name}"
    job_analyst.save_database(db)

    return ApplicationResponse(
        job_id=payload.job_id,
        job_title=job_title,
        company=company,
        cover_letter_markdown=cover_letter_markdown
    )

@app.post("/api/jobs/export-docx")
def export_docx(payload: ExportDocxRequest):
    try:
        import docx
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        title = doc.add_heading(f"Søknad – {payload.job_title}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        p_comp = doc.add_paragraph()
        r_comp = p_comp.add_run(f"{payload.company}")
        r_comp.bold = True
        r_comp.font.size = Pt(12)

        doc.add_paragraph("")

        lines = payload.content_markdown.splitlines()
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
            if line_str.startswith("#"):
                continue
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(line_str)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

        clean_company = "".join(c for c in payload.company if c.isalnum() or c in (" ", "_")).strip().replace(" ", "_")
        file_name = f"Soknad_{clean_company}_Markus.docx"
        file_path = SOKNADSBREV_DIR / file_name
        doc.save(file_path)

        return FileResponse(
            path=str(file_path),
            filename=file_name,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Kunne ikke generere DOCX: {e}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("backend.main:app", host="0.0.0.0", port=8000, reload=True)
